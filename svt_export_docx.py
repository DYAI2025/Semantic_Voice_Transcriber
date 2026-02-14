#!/usr/bin/env python3
"""
SVT Export - DOCX Generation for Therapy Transcriptions
Professional Word document output with speaker separation

Requirements:
    pip install python-docx

Usage:
    from svt_export_docx import generate_docx
    generate_docx(transcript_data, "/output/therapie_bericht.docx")
"""

from datetime import datetime
from typing import List, Dict, Any
from pathlib import Path


def generate_docx(
    transcript_data: List[Dict[str, Any]],
    output_path: str,
    title: str = "Therapie Transkription",
    patient_name: str = "",
    therapist_name: str = "",
    session_date: str = "",
    include_analysis: bool = True,
    analysis_text: str = ""
) -> str:
    """
    Generate a professional DOCX from transcription data.
    
    Args:
        transcript_data: List of transcript segments with speaker/text
        output_path: Path to save DOCX
        title: Document title
        patient_name: Patient name (optional)
        therapist_name: Therapist name (optional)
        session_date: Session date (optional)
        include_analysis: Include analysis section
        analysis_text: Analysis content (optional)
    
    Returns:
        Path to generated DOCX
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    # Create document
    doc = Document()
    
    # Custom styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(24)
    title_style.font.color.rgb = RGBColor(45, 55, 72)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Heading styles
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(18)
    h1_style.font.color.rgb = RGBColor(66, 153, 225)
    
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(14)
    h2_style.font.color.rgb = RGBColor(66, 153, 225)
    
    h3_style = styles['Heading 3']
    h3_style.font.size = Pt(12)
    
    # Body style
    body_style = styles['Normal']
    body_style.font.size = Pt(11)
    body_style.paragraph_format.line_spacing = 1.5
    
    # Title
    title_para = doc.add_paragraph(title, style='Title')
    title_para.paragraph_format.space_after = Pt(20)
    
    # Metadata section
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.style = 'Table Grid'
    
    metadata = []
    if patient_name:
        metadata.append(("Patient:", patient_name))
    if therapist_name:
        metadata.append(("Therapeut:", therapist_name))
    if session_date:
        metadata.append(("Sitzungsdatum:", session_date))
    metadata.append(("Erstellt:", datetime.now().strftime("%d.%m.%Y %H:%M")))
    metadata.append(("Gesamtdauer:", calculate_duration(transcript_data)))
    metadata.append(("Wortzahl:", str(count_words(transcript_data))))
    
    # Fill metadata table
    for i, (label, value) in enumerate(metadata):
        if i == 0:
            row = meta_table.rows[0]
        else:
            row = meta_table.add_row()
        
        cell1 = row.cells[0]
        cell2 = row.cells[1]
        
        cell1.text = label
        cell1.width = Inches(1.5)
        
        cell2.text = value
        cell2.width = Inches(4)
        
        # Style cells
        for cell in [cell1, cell2]:
            paragraph = cell.paragraphs[0]
            paragraph.style = 'Normal'
            paragraph.paragraph_format.space_after = Pt(0)
    
    doc.add_paragraph()
    
    # Summary section
    speaker_summary = get_speaker_summary(transcript_data)
    
    doc.add_heading('📊 Zusammenfassung', level=1)
    
    for spk, data in speaker_summary.items():
        p = doc.add_paragraph()
        p.add_run(f"👤 {spk}: ").bold = True
        p.add_run(f"{data['duration']} ({data['percentage']:.1f}%)")
    
    doc.add_paragraph()
    
    # Divider
    doc.add_paragraph("─" * 60)
    doc.add_paragraph()
    
    # Transcript section
    doc.add_heading('📝 Transkript', level=1)
    
    current_speaker = None
    for seg in transcript_data:
        speaker = seg.get("speaker", "Unknown")
        text = seg.get("text", "")
        start = seg.get("start", 0)
        end = seg.get("end", 0)
        
        # New speaker header
        if speaker != current_speaker:
            doc.add_heading(f"👤 {speaker}", level=2)
            current_speaker = speaker
        
        # Timestamp
        timestamp = f"[{format_time(start)} - {format_time(end)}]"
        p = doc.add_paragraph()
        p.add_run(timestamp).italic = True
        p.add_run(" ")
        
        # Speaker label
        run = p.add_run(f"({speaker}): ")
        run.font.color.rgb = RGBColor(66, 153, 225)
        run.bold = True
        
        # Text
        p.add_run(text)
        p.paragraph_format.space_after = Pt(8)
    
    # Analysis section
    if include_analysis and analysis_text:
        doc.add_page_break()
        doc.add_heading('🧠 Analyse', level=1)
        
        # Parse and add analysis paragraphs
        for paragraph in analysis_text.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip(), style='Normal')
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph("─" * 60)
    doc.add_paragraph(
        "Generiert von SVT Local | Therapie Transkription",
        style='Normal'
    )
    doc.add_paragraph(
        f"Dokument-ID: {generate_document_id()}",
        style='Normal'
    )
    
    # Save
    doc.save(output_path)
    
    return output_path


def create_simple_report(
    transcript_data: List[Dict[str, Any]],
    output_path: str,
    patient_name: str = ""
) -> str:
    """
    Create a simple therapy report DOCX.
    """
    session_date = datetime.now().strftime("%d.%m.%Y")
    
    analysis = generate_simple_analysis(transcript_data)
    
    return generate_docx(
        transcript_data=transcript_data,
        output_path=output_path,
        title="Therapie Transkription",
        patient_name=patient_name,
        session_date=session_date,
        include_analysis=True,
        analysis_text=analysis
    )


def generate_simple_analysis(transcript_data: List[Dict[str, Any]]) -> str:
    """
    Generate a simple analysis summary.
    """
    speaker_times = {}
    total_duration = 0
    
    for seg in transcript_data:
        duration = seg.get("end", 0) - seg.get("start", 0)
        total_duration += duration
        speaker = seg.get("speaker", "Unknown")
        
        if speaker not in speaker_times:
            speaker_times[speaker] = 0
        speaker_times[speaker] += duration
    
    lines = []
    lines.append("Dieses Transkript wurde automatisch erstellt.")
    lines.append("")
    lines.append(f"Gesamtdauer: {int(total_duration // 60)} Minuten")
    lines.append("")
    lines.append("Sprecheraufteilung:")
    
    for spk, duration in sorted(speaker_times.items(), key=lambda x: -x[1]):
        pct = (duration / total_duration * 100) if total_duration > 0 else 0
        lines.append(f"- {spk}: {int(duration // 60)} Min ({pct:.1f}%)")
    
    lines.append("")
    lines.append("Hinweis: Dies ist eine automatische Transkription. Für therapeutische Zwecke empfehlen wir eine manuelle Überprüfung der Inhalte.")
    
    return "\n\n".join(lines)


def calculate_duration(transcript_data: List[Dict[str, Any]]) -> str:
    """Calculate total duration from transcript."""
    total = 0
    for seg in transcript_data:
        total += seg.get("end", 0) - seg.get("start", 0)
    
    hours = int(total // 3600)
    minutes = int((total % 3600) // 60)
    seconds = int(total % 60)
    
    if hours > 0:
        return f"{hours}h {minutes}m {seconds}s"
    return f"{minutes}m {seconds}s"


def count_words(transcript_data: List[Dict[str, Any]]) -> int:
    """Count total words."""
    total = 0
    for seg in transcript_data:
        text = seg.get("text", "")
        total += len(text.split()) if text else 0
    return total


def get_speaker_summary(transcript_data: List[Dict[str, Any]]) -> Dict[str, Dict]:
    """Get speaker statistics."""
    speaker_times = {}
    total_duration = 0
    
    for seg in transcript_data:
        duration = seg.get("end", 0) - seg.get("start", 0)
        total_duration += duration
        speaker = seg.get("speaker", "Unknown")
        
        if speaker not in speaker_times:
            speaker_times[speaker] = {"duration": 0, "percentage": 0}
        speaker_times[speaker]["duration"] += duration
    
    # Calculate percentages
    for spk in speaker_times:
        pct = (speaker_times[spk]["duration"] / total_duration * 100) if total_duration > 0 else 0
        minutes = int(speaker_times[spk]["duration"] // 60)
        seconds = int(speaker_times[spk]["duration"] % 60)
        speaker_times[spk]["duration"] = f"{minutes}m {seconds}s"
        speaker_times[spk]["percentage"] = pct
    
    return speaker_times


def format_time(seconds: float) -> str:
    """Format seconds to MM:SS."""
    minutes = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{minutes:02d}:{secs:02d}"


def generate_document_id() -> str:
    """Generate a unique document ID."""
    import hashlib
    timestamp = str(datetime.now().timestamp()).encode()
    return hashlib.md5(timestamp).hexdigest()[:8].upper()


if __name__ == "__main__":
    # Example usage
    example_data = [
        {"start": 0, "end": 30, "speaker": "Therapeut", "text": "Guten Tag, wie geht es Ihnen heute?"},
        {"start": 30, "end": 90, "speaker": "Patient", "text": "Naja, es geht so. Ich habe schlecht geschlafen."},
        {"start": 90, "end": 150, "speaker": "Therapeut", "text": "Erzählen Sie mir mehr darüber."},
    ]
    
    output = "/tmp/therapie_bericht.docx"
    generate_docx(example_data, output, patient_name="Max Mustermann")
    print(f"DOCX erstellt: {output}")
